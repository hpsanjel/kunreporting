from flask import Flask, render_template, request, redirect, url_for, flash, session, make_response
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from datetime import datetime
import os
from docx import Document
from docx.shared import Inches
import io
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from config import config

app = Flask(__name__)
app.config.from_object(config[os.getenv('FLASK_CONFIG', 'default')])

db = SQLAlchemy(app)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

# Database Models
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password_hash = db.Column(db.String(128))
    full_name = db.Column(db.String(100), nullable=False)
    branch = db.Column(db.String(20), nullable=False)  # 'Oslo' or 'Bergen'
    is_admin = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

class Introduction(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    content = db.Column(db.Text, nullable=False)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    updated_by = db.Column(db.Integer, db.ForeignKey('user.id'))

class Project(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(200), nullable=False)
    description = db.Column(db.Text, nullable=False)
    branch = db.Column(db.String(20), nullable=False)
    created_by = db.Column(db.Integer, db.ForeignKey('user.id'))
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

class GrantApplication(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    project_title = db.Column(db.String(200), nullable=False)
    branch = db.Column(db.String(20), nullable=False)
    applied_by = db.Column(db.Integer, db.ForeignKey('user.id'))
    amount = db.Column(db.Float, nullable=False)
    status = db.Column(db.String(20), nullable=False)  # 'Applied', 'Granted', 'Rejected'
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

class AuditLog(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'))
    action = db.Column(db.String(100), nullable=False)
    table_name = db.Column(db.String(50), nullable=False)
    record_id = db.Column(db.Integer, nullable=False)
    old_values = db.Column(db.Text)
    new_values = db.Column(db.Text)
    timestamp = db.Column(db.DateTime, default=datetime.utcnow)

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# Routes
@app.route('/')
def index():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    return redirect(url_for('login'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        user = User.query.filter_by(username=username).first()
        
        if user and user.check_password(password):
            login_user(user)
            return redirect(url_for('dashboard'))
        else:
            flash('Invalid username or password')
    
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

@app.route('/dashboard')
@login_required
def dashboard():
    projects_count = Project.query.count()
    grants_applied = GrantApplication.query.filter_by(status='Applied').count()
    grants_granted = GrantApplication.query.filter_by(status='Granted').count()
    grants_rejected = GrantApplication.query.filter_by(status='Rejected').count()
    
    return render_template('dashboard.html', 
                         projects_count=projects_count,
                         grants_applied=grants_applied,
                         grants_granted=grants_granted,
                         grants_rejected=grants_rejected)

@app.route('/introduction', methods=['GET', 'POST'])
@login_required
def introduction():
    if not current_user.is_admin:
        flash('Only administrators can edit the introduction.')
        return redirect(url_for('dashboard'))
    
    intro = Introduction.query.first()
    if not intro:
        intro = Introduction(content="Welcome to our company report system.")
        db.session.add(intro)
        db.session.commit()
    
    if request.method == 'POST':
        old_content = intro.content
        intro.content = request.form['content']
        intro.updated_by = current_user.id
        db.session.commit()
        
        # Log the change
        log_audit(current_user.id, 'UPDATE', 'introduction', intro.id, 
                 old_content, intro.content)
        
        flash('Introduction updated successfully!')
        return redirect(url_for('introduction'))
    
    return render_template('introduction.html', intro=intro)

@app.route('/projects')
@login_required
def projects():
    projects = Project.query.order_by(Project.updated_at.desc()).all()
    return render_template('projects.html', projects=projects)

@app.route('/projects/add', methods=['GET', 'POST'])
@login_required
def add_project():
    if request.method == 'POST':
        project = Project(
            title=request.form['title'],
            description=request.form['description'],
            branch=request.form['branch'],
            created_by=current_user.id
        )
        db.session.add(project)
        db.session.commit()
        
        log_audit(current_user.id, 'CREATE', 'project', project.id, 
                 None, f"Title: {project.title}, Branch: {project.branch}")
        
        flash('Project added successfully!')
        return redirect(url_for('projects'))
    
    return render_template('add_project.html')

@app.route('/projects/edit/<int:id>', methods=['GET', 'POST'])
@login_required
def edit_project(id):
    project = Project.query.get_or_404(id)
    
    if request.method == 'POST':
        old_values = f"Title: {project.title}, Description: {project.description}, Branch: {project.branch}"
        
        project.title = request.form['title']
        project.description = request.form['description']
        project.branch = request.form['branch']
        db.session.commit()
        
        new_values = f"Title: {project.title}, Description: {project.description}, Branch: {project.branch}"
        log_audit(current_user.id, 'UPDATE', 'project', project.id, old_values, new_values)
        
        flash('Project updated successfully!')
        return redirect(url_for('projects'))
    
    return render_template('edit_project.html', project=project)

@app.route('/projects/delete/<int:id>')
@login_required
def delete_project(id):
    project = Project.query.get_or_404(id)
    old_values = f"Title: {project.title}, Description: {project.description}, Branch: {project.branch}"
    
    db.session.delete(project)
    db.session.commit()
    
    log_audit(current_user.id, 'DELETE', 'project', id, old_values, None)
    flash('Project deleted successfully!')
    return redirect(url_for('projects'))

@app.route('/grants')
@login_required
def grants():
    grants = GrantApplication.query.order_by(GrantApplication.updated_at.desc()).all()
    return render_template('grants.html', grants=grants)

@app.route('/grants/add', methods=['GET', 'POST'])
@login_required
def add_grant():
    if request.method == 'POST':
        grant = GrantApplication(
            project_title=request.form['project_title'],
            branch=request.form['branch'],
            applied_by=current_user.id,
            amount=float(request.form['amount']),
            status=request.form['status']
        )
        db.session.add(grant)
        db.session.commit()
        
        log_audit(current_user.id, 'CREATE', 'grant', grant.id,
                 None, f"Project: {grant.project_title}, Amount: {grant.amount}, Status: {grant.status}")
        
        flash('Grant application added successfully!')
        return redirect(url_for('grants'))
    
    return render_template('add_grant.html')

@app.route('/grants/edit/<int:id>', methods=['GET', 'POST'])
@login_required
def edit_grant(id):
    grant = GrantApplication.query.get_or_404(id)
    
    if request.method == 'POST':
        old_values = f"Project: {grant.project_title}, Amount: {grant.amount}, Status: {grant.status}"
        
        grant.project_title = request.form['project_title']
        grant.branch = request.form['branch']
        grant.amount = float(request.form['amount'])
        grant.status = request.form['status']
        db.session.commit()
        
        new_values = f"Project: {grant.project_title}, Amount: {grant.amount}, Status: {grant.status}"
        log_audit(current_user.id, 'UPDATE', 'grant', grant.id, old_values, new_values)
        
        flash('Grant application updated successfully!')
        return redirect(url_for('grants'))
    
    return render_template('edit_grant.html', grant=grant)

@app.route('/grants/delete/<int:id>')
@login_required
def delete_grant(id):
    grant = GrantApplication.query.get_or_404(id)
    old_values = f"Project: {grant.project_title}, Amount: {grant.amount}, Status: {grant.status}"
    
    db.session.delete(grant)
    db.session.commit()
    
    log_audit(current_user.id, 'DELETE', 'grant', id, old_values, None)
    flash('Grant application deleted successfully!')
    return redirect(url_for('grants'))

@app.route('/users')
@login_required
def users():
    if not current_user.is_admin:
        flash('Only administrators can manage users.')
        return redirect(url_for('dashboard'))
    
    users = User.query.order_by(User.branch, User.full_name).all()
    return render_template('users.html', users=users)

@app.route('/users/add', methods=['GET', 'POST'])
@login_required
def add_user():
    if not current_user.is_admin:
        flash('Only administrators can create users.')
        return redirect(url_for('dashboard'))
    
    if request.method == 'POST':
        username = request.form['username']
        email = request.form['email']
        full_name = request.form['full_name']
        branch = request.form['branch']
        is_admin = request.form.get('is_admin') == 'on'
        password = request.form['password']
        
        # Check if username already exists
        existing_user = User.query.filter_by(username=username).first()
        if existing_user:
            flash(f'Username "{username}" already exists!')
            return render_template('add_user.html')
        
        # Create new user
        user = User(
            username=username,
            email=email,
            full_name=full_name,
            branch=branch,
            is_admin=is_admin
        )
        user.set_password(password)
        db.session.add(user)
        db.session.commit()
        
        log_audit(current_user.id, 'CREATE', 'user', user.id,
                 None, f"Username: {user.username}, Name: {user.full_name}, Branch: {user.branch}, Role: {'Admin' if user.is_admin else 'Staff'}")
        
        flash(f'User "{username}" created successfully!')
        return redirect(url_for('users'))
    
    return render_template('add_user.html')

@app.route('/users/edit/<int:id>', methods=['GET', 'POST'])
@login_required
def edit_user(id):
    if not current_user.is_admin:
        flash('Only administrators can edit users.')
        return redirect(url_for('dashboard'))
    
    user = User.query.get_or_404(id)
    
    if request.method == 'POST':
        old_values = f"Username: {user.username}, Name: {user.full_name}, Email: {user.email}, Branch: {user.branch}, Role: {'Admin' if user.is_admin else 'Staff'}"
        
        user.username = request.form['username']
        user.email = request.form['email']
        user.full_name = request.form['full_name']
        user.branch = request.form['branch']
        user.is_admin = request.form.get('is_admin') == 'on'
        
        # Update password if provided
        password = request.form.get('password')
        if password:
            user.set_password(password)
        
        db.session.commit()
        
        new_values = f"Username: {user.username}, Name: {user.full_name}, Email: {user.email}, Branch: {user.branch}, Role: {'Admin' if user.is_admin else 'Staff'}"
        log_audit(current_user.id, 'UPDATE', 'user', user.id, old_values, new_values)
        
        flash(f'User "{user.username}" updated successfully!')
        return redirect(url_for('users'))
    
    return render_template('edit_user.html', user=user)

@app.route('/users/delete/<int:id>')
@login_required
def delete_user(id):
    if not current_user.is_admin:
        flash('Only administrators can delete users.')
        return redirect(url_for('dashboard'))
    
    user = User.query.get_or_404(id)
    
    if user.username == 'admin':
        flash('Cannot delete the default admin user!')
        return redirect(url_for('users'))
    
    old_values = f"Username: {user.username}, Name: {user.full_name}, Email: {user.email}, Branch: {user.branch}, Role: {'Admin' if user.is_admin else 'Staff'}"
    
    db.session.delete(user)
    db.session.commit()
    
    log_audit(current_user.id, 'DELETE', 'user', id, old_values, None)
    flash(f'User "{user.username}" deleted successfully!')
    return redirect(url_for('users'))

@app.route('/audit')
@login_required
def audit():
    if not current_user.is_admin:
        flash('Only administrators can view audit logs.')
        return redirect(url_for('dashboard'))
    
    logs = AuditLog.query.order_by(AuditLog.timestamp.desc()).limit(100).all()
    return render_template('audit.html', logs=logs)

@app.route('/report/pdf')
@login_required
def generate_pdf_report():
    intro = Introduction.query.first()
    projects = Project.query.all()
    grants = GrantApplication.query.all()
    
    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    
    # Title
    p.setFont("Helvetica-Bold", 16)
    p.drawString(50, height - 50, "Company Report")
    p.drawString(50, height - 80, f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    
    y_position = height - 120
    
    # Introduction
    if intro:
        p.setFont("Helvetica-Bold", 14)
        p.drawString(50, y_position, "Introduction")
        y_position -= 30
        p.setFont("Helvetica", 10)
        
        intro_lines = intro.content.split('\n')
        for line in intro_lines:
            if y_position < 50:
                p.showPage()
                y_position = height - 50
            p.drawString(50, y_position, line)
            y_position -= 15
        y_position -= 20
    
    # Projects
    p.setFont("Helvetica-Bold", 14)
    p.drawString(50, y_position, "Projects")
    y_position -= 30
    
    for project in projects:
        if y_position < 100:
            p.showPage()
            y_position = height - 50
        
        p.setFont("Helvetica-Bold", 12)
        p.drawString(50, y_position, f"• {project.title}")
        y_position -= 15
        p.setFont("Helvetica", 10)
        p.drawString(70, y_position, f"Branch: {project.branch}")
        y_position -= 12
        desc_lines = project.description.split('\n')
        for line in desc_lines:
            if y_position < 50:
                p.showPage()
                y_position = height - 50
            p.drawString(70, y_position, line)
            y_position -= 12
        y_position -= 10
    
    # Grant Applications
    y_position -= 20
    p.setFont("Helvetica-Bold", 14)
    p.drawString(50, y_position, "Grant Applications")
    y_position -= 30
    
    for grant in grants:
        if y_position < 80:
            p.showPage()
            y_position = height - 50
        
        p.setFont("Helvetica-Bold", 12)
        p.drawString(50, y_position, f"• {grant.project_title}")
        y_position -= 15
        p.setFont("Helvetica", 10)
        p.drawString(70, y_position, f"Branch: {grant.branch}")
        y_position -= 12
        p.drawString(70, y_position, f"Amount: NOK {grant.amount:,.2f}")
        y_position -= 12
        p.drawString(70, y_position, f"Status: {grant.status}")
        y_position -= 20
    
    p.save()
    buffer.seek(0)
    
    response = make_response(buffer.getvalue())
    response.headers['Content-Type'] = 'application/pdf'
    response.headers['Content-Disposition'] = 'attachment; filename=company_report.pdf'
    return response

@app.route('/report/word')
@login_required
def generate_word_report():
    intro = Introduction.query.first()
    projects = Project.query.all()
    grants = GrantApplication.query.all()
    
    doc = Document()
    doc.add_heading('Company Report', 0)
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    
    # Introduction
    if intro:
        doc.add_heading('Introduction', level=1)
        doc.add_paragraph(intro.content)
    
    # Projects
    doc.add_heading('Projects', level=1)
    for project in projects:
        doc.add_heading(project.title, level=2)
        p = doc.add_paragraph(f'Branch: {project.branch}\n')
        p.add_run(project.description)
    
    # Grant Applications
    doc.add_heading('Grant Applications', level=1)
    for grant in grants:
        doc.add_heading(grant.project_title, level=2)
        p = doc.add_paragraph(f'Branch: {grant.branch}\n')
        p.add_run(f'Amount: NOK {grant.amount:,.2f}\n')
        p.add_run(f'Status: {grant.status}')
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    response = make_response(buffer.getvalue())
    response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    response.headers['Content-Disposition'] = 'attachment; filename=company_report.docx'
    return response

def log_audit(user_id, action, table_name, record_id, old_values, new_values):
    log = AuditLog(
        user_id=user_id,
        action=action,
        table_name=table_name,
        record_id=record_id,
        old_values=old_values,
        new_values=new_values
    )
    db.session.add(log)
    db.session.commit()

def create_admin_user():
    admin = User.query.filter_by(username='admin').first()
    if not admin:
        admin = User(
            username='admin',
            email='admin@company.no',
            full_name='System Administrator',
            branch='Oslo',
            is_admin=True
        )
        admin.set_password('admin123')
        db.session.add(admin)
        db.session.commit()
        print("Admin user created: username='admin', password='admin123'")

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
        create_admin_user()
    
    # Use debug mode based on configuration
    debug_mode = app.config.get('DEBUG', False)
    app.run(debug=debug_mode, host='0.0.0.0', port=5001)
